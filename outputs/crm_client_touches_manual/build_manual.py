from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = "/Users/diaoune/automation-agency/outputs/crm_client_touches_manual/client_touches_crm_manual.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(17, 24, 39)
MUTED = RGBColor(71, 85, 105)
FILL_BLUE_GRAY = "E8EEF5"
FILL_NOTE = "F4F6F9"
BORDER = "C9D3DF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])


def set_cell_margins(table, top=80, bottom=80, start=120, end=120):
    tbl_pr = table._tbl.tblPr
    mar = tbl_pr.find(qn("w:tblCellMar"))
    if mar is None:
        mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(mar)
    for name, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_table(table, widths, header=True):
    table.style = "Table Grid"
    set_table_width(table, widths)
    set_cell_margins(table)
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
                    run.font.color.rgb = INK
        if header and r_idx == 0:
            for cell in row.cells:
                set_cell_shading(cell, FILL_BLUE_GRAY)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = DARK_BLUE


def add_para(doc, text="", style=None, bold=False, color=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = color or INK
    run.bold = bold
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(item)
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        r.font.color.rgb = INK


def add_steps(doc, items):
    num_id = new_decimal_numbering(doc)
    for item in items:
        p = doc.add_paragraph()
        p_pr = p._p.get_or_add_pPr()
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            num_pr = OxmlElement("w:numPr")
            p_pr.append(num_pr)
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id_node)
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(item)
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        r.font.color.rgb = INK


def new_decimal_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
        if el.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
        if el.get(qn("w:numId")) is not None
    ]
    abstract_id = (max(abstract_ids) + 1) if abstract_ids else 0
    num_id = (max(num_ids) + 1) if num_ids else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    style_table(table, [9360], header=False)
    cell = table.cell(0, 0)
    set_cell_shading(cell, FILL_NOTE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(10.5)
    r.font.color.rgb = DARK_BLUE
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.name = "Calibri"
    r2.font.size = Pt(10)
    r2.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.line_spacing = 1.25
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(7)
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = BLUE if level in (1, 2) else DARK_BLUE
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)
styles["Normal"].font.color.rgb = INK
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.line_spacing = 1.25
for style_name, size, color in [
    ("Heading 1", 16, BLUE),
    ("Heading 2", 13, BLUE),
    ("Heading 3", 12, DARK_BLUE),
]:
    style = styles[style_name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = color

header = section.header.paragraphs[0]
header.text = "Client Touch CRM Manual"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header.runs[0].font.name = "Calibri"
header.runs[0].font.size = Pt(9)
header.runs[0].font.color.rgb = MUTED

footer = section.footer.paragraphs[0]
footer.text = "Use with: client_touches_crm.xlsx"
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer.runs[0].font.name = "Calibri"
footer.runs[0].font.size = Pt(9)
footer.runs[0].font.color.rgb = MUTED

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(8)
title.paragraph_format.space_after = Pt(4)
r = title.add_run("Client Touch CRM User Manual")
r.font.name = "Calibri"
r.font.size = Pt(23)
r.font.bold = True
r.font.color.rgb = RGBColor(0, 0, 0)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(14)
r = subtitle.add_run("How to track outreach, responses, and follow-ups in your CRM spreadsheet")
r.font.name = "Calibri"
r.font.size = Pt(12)
r.font.color.rgb = MUTED

meta = doc.add_table(rows=4, cols=2)
rows = [
    ("Workbook", "Client Touches CRM.xlsx"),
    ("Primary use", "Track LinkedIn/client touches, replies, next follow-up dates, and daily actions."),
    ("Starting context", "Outreach started June 7, 2026. Michelle Mgbo is currently the only respondent."),
    ("Recommended rhythm", "Update the spreadsheet immediately after each touch; review the Dashboard once per workday."),
]
for i, (label, value) in enumerate(rows):
    meta.cell(i, 0).text = label
    meta.cell(i, 1).text = value
style_table(meta, [2700, 6660], header=False)
for row in meta.rows:
    set_cell_shading(row.cells[0], FILL_BLUE_GRAY)
    for p in row.cells[0].paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = DARK_BLUE

add_heading(doc, "Quick Start", 1)
add_steps(doc, [
    "Open the Dashboard tab first to see the current reply count, pending follow-ups, and action queue.",
    "When you message someone, add or update the person in Contacts and log the exact activity in Touch Log.",
    "Use Next Follow-up Date to decide when to reach out again. Keep the date realistic so the Dashboard queue stays useful.",
    "When someone replies, change their Status to Responded, add the Response Date, and set Next Action to a concrete step such as Reply / book next step.",
    "At the end of each day, scan the Dashboard for overdue or due-today follow-ups and update any stale records.",
])

add_note(
    doc,
    "Operating rule",
    "The Contacts tab is the current state of each prospect. The Touch Log is the history of what happened. Update both when a meaningful touch occurs.",
)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
add_heading(doc, "What Each Tab Does", 1)
tab_table = doc.add_table(rows=1, cols=3)
tab_table.rows[0].cells[0].text = "Tab"
tab_table.rows[0].cells[1].text = "Use It For"
tab_table.rows[0].cells[2].text = "Update Frequency"
tab_rows = [
    ("Dashboard", "Daily view of contacted prospects, response rate, pending follow-ups, and the action queue.", "Review daily; formulas update automatically."),
    ("Contacts", "One row per prospect. This is where status, next follow-up date, priority, notes, and next action live.", "Update after every meaningful change."),
    ("Touch Log", "Chronological record of every outreach message, reply, call, or follow-up.", "Add a row after each touch."),
    ("Follow-up Templates", "Reusable message starters for replies, no-response follow-ups, booking, and nurture situations.", "Edit when your messaging improves."),
    ("Lists", "Dropdown/source values for statuses, priorities, channels, and touch types.", "Edit only when you want to add new CRM categories."),
]
for row in tab_rows:
    cells = tab_table.add_row().cells
    for idx, text in enumerate(row):
        cells[idx].text = text
style_table(tab_table, [1800, 5460, 2100])

add_heading(doc, "Daily Workflow", 1)
add_heading(doc, "Morning Review", 2)
add_bullets(doc, [
    "Open Dashboard and look at Follow-ups due today and Overdue follow-ups.",
    "Start with any Responded contacts, especially Michelle, because active conversations decay faster than cold follow-ups.",
    "Use the Action Queue to decide what must happen today.",
])
add_heading(doc, "After Sending a Message", 2)
add_steps(doc, [
    "Go to Touch Log and add a new row with the date, contact, channel, touch type, direction, outcome, message notes, and next step.",
    "Go to Contacts and update Last Touch Date, Last Message Snippet, Notes, and Next Follow-up Date.",
    "If the message is a follow-up to a non-responder, keep Status as Pending Follow-up until they reply or opt out.",
])
add_heading(doc, "After Receiving a Reply", 2)
add_steps(doc, [
    "Change Status in Contacts to Responded.",
    "Enter the Response Date.",
    "Set Priority to High if the reply suggests real interest or a near-term pain point.",
    "Write a clear Next Action such as Reply / book next step, Send example workflow, or Ask discovery question.",
    "Log the reply in Touch Log so the conversation history stays complete.",
])

add_heading(doc, "Field Reference", 1)
field_table = doc.add_table(rows=1, cols=3)
field_table.rows[0].cells[0].text = "Field"
field_table.rows[0].cells[1].text = "Meaning"
field_table.rows[0].cells[2].text = "How to Use"
field_rows = [
    ("Status", "Current relationship state for the prospect.", "Use Pending Follow-up, Responded, Meeting Booked, Nurture, Not Interested, Do Not Contact, or Closed Won."),
    ("Next Follow-up Date", "The next date you should take action.", "Set this every time you send a message or get a reply. This drives the Dashboard."),
    ("Priority", "How important the next action is.", "Use High for active replies or strong-fit prospects; Medium for normal follow-ups; Low for nurture."),
    ("Notes", "Context you need before writing the next message.", "Capture pain points, objections, relevant company details, or what to personalize."),
    ("Last Message Snippet", "Short reminder of the last outbound or inbound message.", "Keep it brief enough to scan from the Contacts tab."),
    ("Next Action", "The next move you should make.", "Make this action-oriented, not vague. Example: Send Michelle booking question."),
]
for row in field_rows:
    cells = field_table.add_row().cells
    for idx, text in enumerate(row):
        cells[idx].text = text
style_table(field_table, [2100, 3300, 3960])

add_heading(doc, "Using the Follow-up Templates", 1)
add_para(doc, "The templates are starting points. Personalize the first sentence whenever possible so your outreach does not feel copied and pasted.")
add_bullets(doc, [
    "Use Responder reply for Michelle and any future prospect who replies.",
    "Use No-response follow-up 1 about two to three days after the first message if there is no reply.",
    "Use No-response follow-up 2 about five to seven days after the first follow-up if the prospect still has not replied.",
    "Use Book meeting only after a prospect shows interest or asks to learn more.",
    "Use Not now when someone gives a soft no but leaves timing open.",
])

add_heading(doc, "Status Rules", 1)
rules = doc.add_table(rows=1, cols=3)
rules.rows[0].cells[0].text = "Situation"
rules.rows[0].cells[1].text = "Set Status To"
rules.rows[0].cells[2].text = "Next Action"
rule_rows = [
    ("Sent initial or follow-up message; no reply yet", "Pending Follow-up", "Set the next follow-up date."),
    ("Prospect replies with any meaningful response", "Responded", "Reply with a discovery question or meeting ask."),
    ("Prospect agrees to a call", "Meeting Booked", "Add meeting details to Notes and Touch Log."),
    ("Prospect says later / not right now", "Nurture", "Set a future follow-up date."),
    ("Prospect says no", "Not Interested", "Stop active outreach unless they invite future contact."),
    ("Prospect asks not to be contacted", "Do Not Contact", "Do not message again."),
]
for row in rule_rows:
    cells = rules.add_row().cells
    for idx, text in enumerate(row):
        cells[idx].text = text
style_table(rules, [3300, 2100, 3960])

add_heading(doc, "Recommended Cadence", 1)
add_bullets(doc, [
    "Day 0: Send the first message after connecting.",
    "Day 2 or 3: Send No-response follow-up 1 if there is no reply.",
    "Day 5 to 7: Send No-response follow-up 2 if there is still no reply.",
    "After a reply: respond within one business day and aim to move toward a discovery question or short meeting.",
])

add_heading(doc, "Quality Checks", 1)
add_bullets(doc, [
    "Every active prospect should have a Status and Next Follow-up Date.",
    "Every sent or received message should have a Touch Log entry.",
    "Dashboard response rate should make sense based on the Contacts tab.",
    "No one marked Do Not Contact should have a future follow-up date.",
    "The Next Action column should say what to do next, not just repeat the status.",
])

add_note(
    doc,
    "Michelle-specific next move",
    "Because Michelle is the only current respondent, prioritize a tailored reply before sending more cold follow-ups. Use the Responder reply template, then adapt it to what she actually said.",
)

doc.save(OUT)
print(OUT)
